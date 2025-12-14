"""
PARSER DE DOCUMENTOS
====================

Procesa y extrae texto de múltiples formatos de documento:
- PDF
- DOCX (Word)
- TXT
- CSV
- JSON
- Markdown

Dependencias:
    pip install PyPDF2 python-docx chardet

Autor: Litper IA System
Versión: 1.0.0
"""

import json
import csv
import io
import asyncio
from typing import Dict
from pathlib import Path
from loguru import logger

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    logger.warning("PyPDF2 no instalado, PDFs no soportados")
    PYPDF2_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    logger.warning("python-docx no instalado, DOCX no soportados")
    DOCX_AVAILABLE = False

try:
    import chardet
    CHARDET_AVAILABLE = True
except ImportError:
    logger.debug("chardet no instalado, usando encoding por defecto")
    CHARDET_AVAILABLE = False


class DocumentParser:
    """
    Parser universal de documentos.

    Soporta múltiples formatos y extrae texto limpio
    para incorporar a la base de conocimiento.

    Ejemplo de uso:
        parser = DocumentParser()

        # Desde archivo
        texto = await parser.parsear("/ruta/archivo.pdf")

        # Desde bytes
        texto = await parser.parsear_bytes(
            contenido=bytes_del_archivo,
            nombre="documento.docx"
        )

    Tip: Para PDFs escaneados (imágenes), considera usar OCR.
    Este parser solo extrae texto embebido.
    """

    def __init__(self):
        """Inicializa el parser con configuraciones por defecto."""
        self.max_file_size = 50 * 1024 * 1024  # 50 MB
        self.encodings_fallback = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

    async def parsear(self, ruta_archivo: str) -> str:
        """
        Parsea un archivo desde una ruta.

        Args:
            ruta_archivo: Ruta al archivo a procesar

        Returns:
            Texto extraído del documento

        Tip: El tipo de archivo se detecta automáticamente
        por la extensión.
        """

        ruta = Path(ruta_archivo)

        if not ruta.exists():
            raise FileNotFoundError(f"Archivo no encontrado: {ruta_archivo}")

        if ruta.stat().st_size > self.max_file_size:
            raise ValueError(f"Archivo muy grande (max {self.max_file_size // 1024 // 1024}MB)")

        extension = ruta.suffix.lower()
        logger.info(f"📄 Parseando archivo {extension}: {ruta.name}")

        # Leer contenido
        with open(ruta, 'rb') as f:
            contenido = f.read()

        return await self.parsear_bytes(contenido, ruta.name)

    async def parsear_archivo(self, ruta_archivo: str) -> str:
        """Alias de parsear() para compatibilidad."""
        return await self.parsear(ruta_archivo)

    async def parsear_bytes(
        self,
        contenido: bytes,
        nombre: str
    ) -> str:
        """
        Parsea contenido desde bytes.

        Args:
            contenido: Bytes del archivo
            nombre: Nombre del archivo (para detectar tipo)

        Returns:
            Texto extraído del documento

        Tip: Útil cuando recibes archivos desde uploads HTTP.
        """

        extension = Path(nombre).suffix.lower()

        # Selector de parser según extensión
        parsers = {
            '.pdf': self._parsear_pdf,
            '.docx': self._parsear_docx,
            '.doc': self._parsear_docx,  # Intentar como docx
            '.txt': self._parsear_texto,
            '.md': self._parsear_texto,
            '.markdown': self._parsear_texto,
            '.csv': self._parsear_csv,
            '.json': self._parsear_json,
            '.xml': self._parsear_xml,
            '.html': self._parsear_html,
            '.htm': self._parsear_html,
        }

        parser = parsers.get(extension)

        if not parser:
            # Intentar como texto plano
            logger.warning(f"Extensión {extension} no reconocida, intentando como texto")
            return await self._parsear_texto(contenido)

        return await parser(contenido)

    async def _parsear_pdf(self, contenido: bytes) -> str:
        """
        Extrae texto de un PDF.

        Proceso:
        1. Lee el PDF página por página
        2. Extrae texto de cada página
        3. Concatena todo con separadores

        Limitación: No extrae texto de imágenes (necesita OCR).

        Tip: Para PDFs complejos con tablas, considera
        usar herramientas especializadas como Tabula.
        """

        if not PYPDF2_AVAILABLE:
            raise ImportError("PyPDF2 no está instalado. Instala con: pip install PyPDF2")

        try:
            # Crear reader desde bytes
            pdf_file = io.BytesIO(contenido)
            reader = PyPDF2.PdfReader(pdf_file)

            textos = []
            total_paginas = len(reader.pages)

            logger.debug(f"PDF tiene {total_paginas} páginas")

            for i, pagina in enumerate(reader.pages):
                try:
                    texto = pagina.extract_text()
                    if texto and texto.strip():
                        textos.append(f"[Página {i + 1}]")
                        textos.append(texto.strip())
                        textos.append("")  # Línea vacía entre páginas
                except Exception as e:
                    logger.warning(f"Error en página {i + 1}: {e}")
                    continue

            if not textos:
                raise ValueError(
                    "No se pudo extraer texto del PDF. "
                    "Puede ser un PDF escaneado (imágenes) que requiere OCR."
                )

            texto_final = '\n'.join(textos)
            logger.success(f"✅ PDF parseado ({len(texto_final)} chars, {total_paginas} páginas)")

            return texto_final

        except PyPDF2.errors.PdfReadError as e:
            raise ValueError(f"PDF corrupto o protegido: {e}")

    async def _parsear_docx(self, contenido: bytes) -> str:
        """
        Extrae texto de un documento Word (DOCX).

        Extrae:
        - Párrafos
        - Texto de tablas
        - Encabezados

        Tip: Las imágenes y gráficos no se extraen.
        """

        if not DOCX_AVAILABLE:
            raise ImportError("python-docx no está instalado. Instala con: pip install python-docx")

        try:
            doc_file = io.BytesIO(contenido)
            doc = DocxDocument(doc_file)

            textos = []

            # Extraer párrafos
            for para in doc.paragraphs:
                texto = para.text.strip()
                if texto:
                    textos.append(texto)

            # Extraer texto de tablas
            for tabla in doc.tables:
                for fila in tabla.rows:
                    celdas = [celda.text.strip() for celda in fila.cells]
                    if any(celdas):
                        textos.append(' | '.join(celdas))

            if not textos:
                raise ValueError("El documento DOCX está vacío")

            texto_final = '\n\n'.join(textos)
            logger.success(f"✅ DOCX parseado ({len(texto_final)} chars)")

            return texto_final

        except Exception as e:
            raise ValueError(f"Error procesando DOCX: {e}")

    async def _parsear_texto(self, contenido: bytes) -> str:
        """
        Parsea texto plano con detección de encoding.

        Intenta múltiples encodings si el primero falla.

        Tip: Archivos Markdown (.md) también se procesan aquí.
        """

        # Detectar encoding
        encoding = 'utf-8'

        if CHARDET_AVAILABLE:
            detected = chardet.detect(contenido)
            if detected and detected.get('encoding'):
                encoding = detected['encoding']
                logger.debug(f"Encoding detectado: {encoding}")

        # Intentar decodificar
        for enc in [encoding] + self.encodings_fallback:
            try:
                texto = contenido.decode(enc)
                logger.success(f"✅ Texto parseado con {enc}")
                return texto
            except UnicodeDecodeError:
                continue

        raise ValueError("No se pudo decodificar el archivo de texto")

    async def _parsear_csv(self, contenido: bytes) -> str:
        """
        Convierte CSV a texto tabular legible.

        Formato de salida:
            Columna1 | Columna2 | Columna3
            Valor1   | Valor2   | Valor3
            ...

        Tip: Para CSVs grandes, considera cargar solo
        las primeras N filas.
        """

        try:
            # Decodificar
            texto = contenido.decode('utf-8', errors='replace')

            # Parsear CSV
            reader = csv.reader(io.StringIO(texto))
            filas = list(reader)

            if not filas:
                raise ValueError("CSV vacío")

            # Formatear como tabla
            lineas = []

            # Headers
            if len(filas) > 0:
                lineas.append("ENCABEZADOS: " + ' | '.join(filas[0]))
                lineas.append("-" * 50)

            # Datos (máximo 1000 filas)
            for fila in filas[1:1001]:
                lineas.append(' | '.join(fila))

            if len(filas) > 1001:
                lineas.append(f"... y {len(filas) - 1001} filas más")

            texto_final = '\n'.join(lineas)
            logger.success(f"✅ CSV parseado ({len(filas)} filas)")

            return texto_final

        except Exception as e:
            raise ValueError(f"Error procesando CSV: {e}")

    async def _parsear_json(self, contenido: bytes) -> str:
        """
        Convierte JSON a texto legible.

        Formatea el JSON con indentación para fácil lectura.

        Tip: Útil para cargar configuraciones o datos estructurados.
        """

        try:
            texto = contenido.decode('utf-8')
            datos = json.loads(texto)

            # Formatear JSON legible
            texto_formateado = json.dumps(
                datos,
                indent=2,
                ensure_ascii=False
            )

            logger.success("✅ JSON parseado")
            return texto_formateado

        except json.JSONDecodeError as e:
            raise ValueError(f"JSON inválido: {e}")

    async def _parsear_xml(self, contenido: bytes) -> str:
        """
        Extrae texto de XML.

        Elimina tags y conserva solo el contenido textual.
        """

        try:
            import re

            texto = contenido.decode('utf-8', errors='replace')

            # Eliminar tags XML
            texto_limpio = re.sub(r'<[^>]+>', ' ', texto)

            # Limpiar espacios
            texto_limpio = re.sub(r'\s+', ' ', texto_limpio).strip()

            logger.success("✅ XML parseado")
            return texto_limpio

        except Exception as e:
            raise ValueError(f"Error procesando XML: {e}")

    async def _parsear_html(self, contenido: bytes) -> str:
        """
        Extrae texto de HTML.

        Usa BeautifulSoup si está disponible, sino regex.

        Tip: Para páginas web complejas, usa WebScraper en su lugar.
        """

        try:
            from bs4 import BeautifulSoup

            html = contenido.decode('utf-8', errors='replace')
            soup = BeautifulSoup(html, 'html.parser')

            # Eliminar scripts y estilos
            for script in soup(["script", "style"]):
                script.decompose()

            texto = soup.get_text(separator='\n', strip=True)
            logger.success("✅ HTML parseado")
            return texto

        except ImportError:
            # Fallback sin BeautifulSoup
            import re

            texto = contenido.decode('utf-8', errors='replace')
            texto = re.sub(r'<script[^>]*>.*?</script>', '', texto, flags=re.DOTALL)
            texto = re.sub(r'<style[^>]*>.*?</style>', '', texto, flags=re.DOTALL)
            texto = re.sub(r'<[^>]+>', ' ', texto)
            texto = re.sub(r'\s+', ' ', texto).strip()

            return texto

    def get_formatos_soportados(self) -> Dict[str, str]:
        """
        Retorna lista de formatos soportados con su estado.

        Returns:
            {
                '.pdf': 'disponible' o 'requiere PyPDF2',
                '.docx': 'disponible' o 'requiere python-docx',
                ...
            }

        Tip: Usa esto para mostrar al usuario qué archivos puede cargar.
        """

        return {
            '.pdf': 'disponible' if PYPDF2_AVAILABLE else 'requiere PyPDF2',
            '.docx': 'disponible' if DOCX_AVAILABLE else 'requiere python-docx',
            '.doc': 'limitado' if DOCX_AVAILABLE else 'requiere python-docx',
            '.txt': 'disponible',
            '.md': 'disponible',
            '.csv': 'disponible',
            '.json': 'disponible',
            '.xml': 'disponible',
            '.html': 'disponible'
        }


# ==================== PARA TESTING ====================

if __name__ == "__main__":
    async def test():
        parser = DocumentParser()

        print("Formatos soportados:")
        for ext, estado in parser.get_formatos_soportados().items():
            print(f"  {ext}: {estado}")

        # Test con archivo de texto
        print("\nTesteando parser de texto...")
        texto = await parser.parsear_bytes(
            b"Este es un texto de prueba.\nCon varias lineas.",
            "test.txt"
        )
        print(f"Texto extraído: {texto}")

    asyncio.run(test())
